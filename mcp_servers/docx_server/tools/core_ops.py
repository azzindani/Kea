import docx
import structlog
import os
from typing import Any, Dict, List, Optional, Union
from contextlib import contextmanager

logger = structlog.get_logger()

def create_document() -> str:
    """Create a new blank document and return its default temporary path or just a success message."""
    # In a stateless server, we usually manipulate files on disk. 
    # This tool might just return a confirmation, but actual paths are handled by the user providing them.
    # For now, we'll assume the user provides output paths for save.
    return "Ready to create. Use save_document to finalize."

def open_document(path: str) -> docx.document.Document:
    """Helper to open a document."""
    try:
        if not os.path.exists(path):
             # If strictly opening, error. If we want auto-create, that's different.
             # Let's align with plan: explicit create/open.
             raise FileNotFoundError(f"Document not found: {path}")
        return docx.Document(path)
    except Exception as e:
        logger.error("failed_to_open_docx", path=path, error=str(e))
        raise

def save_document(doc: docx.document.Document, path: str) -> str:
    """Helper to save document."""
    try:
        doc.save(path)
        return f"Document saved to {path}"
    except Exception as e:
        logger.error("failed_to_save_docx", path=path, error=str(e))
        raise

def validate_docx(path: str) -> Dict[str, Any]:
    """Check if file is valid docx."""
    try:
        doc = docx.Document(path)
        # Check some basic property to ensure it loaded
        paras = len(doc.paragraphs)
        return {"valid": True, "paragraphs": paras, "tables": len(doc.tables)}
    except Exception as e:
        return {"valid": False, "error": str(e)}

# Property Ops (Basic)
def get_document_properties(path: str) -> Dict[str, Any]:
    """Get core properties."""
    doc = open_document(path)
    props = doc.core_properties
    return {
        "author": props.author,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "title": props.title,
        "revision": props.revision
    }

def set_document_properties(path: str, author: Optional[str] = None, title: Optional[str] = None) -> str:
    """Set core properties."""
    doc = open_document(path)
    if author: doc.core_properties.author = author
    if title: doc.core_properties.title = title
    doc.save(path)
    return "Properties updated."
