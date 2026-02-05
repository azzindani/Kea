from mcp_servers.docx_server.tools.core_ops import open_document, save_document
import docx
from typing import List, Dict, Any, Optional

def add_paragraph(path: str, text: str, style: Optional[str] = None) -> str:
    """Add a new paragraph to the end of the document."""
    doc = open_document(path)
    doc.add_paragraph(text, style=style)
    save_document(doc, path)
    return "Paragraph added."

def add_heading(path: str, text: str, level: int = 1) -> str:
    """Add a heading (level 0-9)."""
    doc = open_document(path)
    if not (0 <= level <= 9): level = 1
    doc.add_heading(text, level=level)
    save_document(doc, path)
    return f"Heading level {level} added."

def add_run(path: str, paragraph_index: int, text: str, bold: bool = False, italic: bool = False, color_rgb: Optional[List[int]] = None) -> str:
    """Add a stylized run to a specific paragraph."""
    doc = open_document(path)
    if paragraph_index >= len(doc.paragraphs):
        raise IndexError("Paragraph index out of range")
    para = doc.paragraphs[paragraph_index]
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color_rgb and len(color_rgb) == 3:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*color_rgb)
    save_document(doc, path)
    return "Run added."

def read_paragraphs(path: str) -> List[str]:
    """Read all paragraphs text."""
    doc = open_document(path)
    return [p.text for p in doc.paragraphs]

def read_paragraph_at_index(path: str, index: int) -> str:
    """Read text of specific paragraph."""
    doc = open_document(path)
    if index >= len(doc.paragraphs):
        return ""
    return doc.paragraphs[index].text

def insert_paragraph_before(path: str, index: int, text: str, style: Optional[str] = None) -> str:
    """Insert text before a specific paragraph."""
    doc = open_document(path)
    if index >= len(doc.paragraphs):
        doc.add_paragraph(text, style=style) # Fallback to append
    else:
        para = doc.paragraphs[index]
        para.insert_paragraph_before(text, style=style)
    save_document(doc, path)
    return "Paragraph inserted."

def delete_paragraph(path: str, index: int) -> str:
    """Delete a paragraph by index."""
    doc = open_document(path)
    if index >= len(doc.paragraphs):
         return "Index out of range."
    p = doc.paragraphs[index]
    p._element.getparent().remove(p._element)
    save_document(doc, path)
    return "Paragraph deleted."

def replace_text(path: str, old_text: str, new_text: str) -> int:
    """Find and replace text in all paragraphs."""
    doc = open_document(path)
    count = 0
    for p in doc.paragraphs:
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text) # Simple replacement loses run formatting
            count += 1
    save_document(doc, path)
    return count

def get_text_stats(path: str) -> Dict[str, int]:
    """Count words, chars, paragraphs."""
    doc = open_document(path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return {
        "paragraphs": len(doc.paragraphs),
        "characters": len(text),
        "words": len(text.split())
    }

def clear_document_content(path: str) -> str:
    """Remove all body content (paragraphs and tables)."""
    # This is complex in xml, simpler approach is to delete all P and Tbl elements in body
    doc = open_document(path)
    body = doc._body
    body.clear_content() # python-docx internal method if available or iterate element
    # doc._element.body.clear() is safer for pure xml clear
    doc._element.body.clear()
    save_document(doc, path)
    return "Content cleared."
