from mcp_servers.docx_server.tools.core_ops import open_document, save_document, create_document
import docx
from typing import Dict, Any, List
import json

def inspect_document_structure(path: str) -> Dict[str, Any]:
    """Return hierarchy of document content."""
    doc = open_document(path)
    structure = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "inline_shapes": len(doc.inline_shapes),
        "preview": [p.text[:50] + "..." for p in doc.paragraphs[:5]]
    }
    return structure

def create_invoice(path: str, data: Dict[str, Any]) -> str:
    """Generate a formatted invoice from JSON data."""
    doc = docx.Document()
    doc.add_heading('INVOICE', 0)
    
    doc.add_paragraph(f"Invoice Number: {data.get('invoice_number', 'N/A')}")
    doc.add_paragraph(f"Date: {data.get('date', 'N/A')}")
    doc.add_paragraph(f"Bill To: {data.get('bill_to', 'N/A')}")
    
    doc.add_heading('Items', level=2)
    
    items = data.get('items', [])
    if items:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Description'
        hdr[1].text = 'Quantity'
        hdr[2].text = 'Price'
        
        total = 0.0
        for item in items:
            row = table.add_row().cells
            row[0].text = item.get('description', '')
            qty = item.get('quantity', 0)
            price = item.get('price', 0.0)
            row[1].text = str(qty)
            row[2].text = f"${price:.2f}"
            total += (qty * price)
            
        doc.add_paragraph(f"\nTotal: ${total:.2f}", style='Heading 3')
    
    doc.save(path)
    return f"Invoice generated at {path}"

def convert_to_markdown(path: str) -> str:
    """Convert document to Markdown (heuristic)."""
    doc = open_document(path)
    md_lines = []
    for p in doc.paragraphs:
        style = p.style.name
        text = p.text
        if 'Heading 1' in style: md_lines.append(f"# {text}")
        elif 'Heading 2' in style: md_lines.append(f"## {text}")
        elif 'Heading 3' in style: md_lines.append(f"### {text}")
        elif 'List Bullet' in style: md_lines.append(f"- {text}")
        elif 'List Number' in style: md_lines.append(f"1. {text}") # Simple numbering
        else: md_lines.append(text)
        md_lines.append("") # Spacing
    
    # Tables
    for t in doc.tables:
        if t.rows:
            # Header
            headers = [c.text for c in t.rows[0].cells]
            md_lines.append("| " + " | ".join(headers) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in t.rows[1:]:
                cells = [c.text for c in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")
            
    return "\n".join(md_lines)

def template_fill(template_path: str, output_path: str, replacements: Dict[str, str]) -> str:
    """Fill a template document replacing {{keys}} with values."""
    doc = open_document(template_path)
    
    def replace_in_p(p):
        for key, val in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                # Naive replacement doesn't handle run-splitting well, but functional for simple cases
                p.text = p.text.replace(placeholder, str(val))
    
    for p in doc.paragraphs:
        replace_in_p(p)
        
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_p(p)
                    
    doc.save(output_path)
    return f"Template filled and saved to {output_path}"
