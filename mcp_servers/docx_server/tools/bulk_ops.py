from mcp_servers.docx_server.tools.core_ops import open_document, save_document
import docx
from typing import Dict, Any, List

def bulk_read_text(path: str) -> str:
    """Read full text of document (paragraphs + tables)."""
    doc = open_document(path)
    text = []
    for element in doc.element.body:
        if isinstance(element, docx.oxml.text.paragraph.CT_P):
             # This is low level, safer to iterate object lists in order if possible.
             # python-docx doesn't easily interleave.
             pass
    # Simple approach: Join paras then tables? or try to respect order?
    # Respected order is hard in python-docx without xml iteration.
    # We will do: Paras, then Tables (standard extraction)
    # OR iterate body elements via _body
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            full_text.append(row_text)
    return "\n".join(full_text)

def bulk_read_tables(path: str) -> List[List[List[str]]]:
    """Get all tables."""
    # same as read_all_tables
    doc = open_document(path)
    return [[[cell.text for cell in row.cells] for row in t.rows] for t in doc.tables]

def bulk_replace_text(path: str, replacements: Dict[str, str]) -> int:
    """Batch replace multiple text patterns."""
    doc = open_document(path)
    count = 0
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
                count += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                     for old, new in replacements.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)
                            count += 1
    save_document(doc, path)
    return count

def bulk_add_paragraphs(path: str, paragraphs: List[str], style: str = "Normal") -> str:
    """Add multiple paragraphs at once."""
    doc = open_document(path)
    for text in paragraphs:
        doc.add_paragraph(text, style=style)
    save_document(doc, path)
    return f"Added {len(paragraphs)} paragraphs."

def bulk_add_table_data(path: str, table_index: int, data: List[List[str]]) -> str:
    """Populate table from 2D list."""
    doc = open_document(path)
    if table_index >= len(doc.tables): raise IndexError("Table index out of range")
    table = doc.tables[table_index]
    
    # Check dimensions
    needed_rows = len(data)
    current_rows = len(table.rows)
    # Add rows if needed
    for _ in range(needed_rows - current_rows):
        table.add_row()
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                row.cells[j].text = str(cell_text)
    save_document(doc, path)
    return "Table populated."

def extract_all_images(path: str, output_dir: str) -> List[str]:
    """Save all embedded images to directory."""
    doc = open_document(path)
    import os
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    saved_files = []
    
    # Relationships are where images live
    # Iterating doc.part.rels
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
             filename = os.path.basename(rel.target_part.partname)
             out_path = os.path.join(output_dir, filename)
             with open(out_path, "wb") as f:
                 f.write(rel.target_part.blob)
             saved_files.append(out_path)
    return saved_files
