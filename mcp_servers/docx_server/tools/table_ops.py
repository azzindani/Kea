from mcp_servers.docx_server.tools.core_ops import open_document, save_document
import docx
from typing import List, Dict, Any, Optional
import pandas as pd

def add_table(path: str, rows: int, cols: int, style: str = "Table Grid") -> str:
    """Add a new table."""
    doc = open_document(path)
    doc.add_table(rows=rows, cols=cols, style=style)
    save_document(doc, path)
    return "Table added."

def _get_table(doc: Any, table_index: int) -> Any:
    if table_index >= len(doc.tables): raise IndexError("Table index out of range")
    return doc.tables[table_index]

def add_row(path: str, table_index: int, cells_text: List[str]) -> str:
    """Add a row to a table with content."""
    doc = open_document(path)
    table = _get_table(doc, table_index)
    row = table.add_row()
    for i, text in enumerate(cells_text):
        if i < len(row.cells):
            row.cells[i].text = str(text)
    save_document(doc, path)
    return "Row added."

def add_column(path: str, table_index: int, width_inches: float = 1.0) -> str:
    """Add a column to a table."""
    doc = open_document(path)
    table = _get_table(doc, table_index)
    table.add_column(width=docx.shared.Inches(width_inches))
    # python-docx add_column is tricky, it doesn't auto-populate cells easily, 
    # but the API exists. It adds cells to all rows.
    save_document(doc, path)
    return "Column added."

def set_cell_text(path: str, table_index: int, row: int, col: int, text: str) -> str:
    """Set specific cell text."""
    doc = open_document(path)
    table = _get_table(doc, table_index)
    try:
        table.cell(row, col).text = text
    except IndexError:
        return "Cell out of range."
    save_document(doc, path)
    return "Cell updated."

def read_table(path: str, table_index: int) -> List[List[str]]:
    """Read table as 2D list."""
    doc = open_document(path)
    if table_index >= len(doc.tables): return []
    table = doc.tables[table_index]
    return [[cell.text for cell in row.cells] for row in table.rows]

def read_all_tables(path: str) -> List[List[List[str]]]:
    """Read all tables."""
    doc = open_document(path)
    return [[[cell.text for cell in row.cells] for row in t.rows] for t in doc.tables]

def read_table_df(path: str, table_index: int) -> List[Dict[str, Any]]:
    """Read table as DataFrame records (assumes row 0 is header)."""
    data = read_table(path, table_index)
    if not data: return []
    if len(data) < 2: return [] # Need header and data
    df = pd.DataFrame(data[1:], columns=data[0])
    return df.to_dict(orient='records')

def style_table(path: str, table_index: int, style_name: str) -> str:
    """Apply style to table."""
    doc = open_document(path)
    table = _get_table(doc, table_index)
    table.style = style_name
    save_document(doc, path)
    return "Style applied."

def merge_cells(path: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """Merge a range of cells (rectangular area)."""
    doc = open_document(path)
    table = _get_table(doc, table_index)
    a = table.cell(start_row, start_col)
    b = table.cell(end_row, end_col)
    a.merge(b)
    save_document(doc, path)
    return "Cells merged."
