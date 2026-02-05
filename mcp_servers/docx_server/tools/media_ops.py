from mcp_servers.docx_server.tools.core_ops import open_document, save_document
import docx
import io
import base64
from docx.shared import Inches, Cm
from typing import Optional

def add_image(path: str, image_path: str, width_inches: Optional[float] = None, height_inches: Optional[float] = None) -> str:
    """Add an image from a file path."""
    doc = open_document(path)
    w = Inches(width_inches) if width_inches else None
    h = Inches(height_inches) if height_inches else None
    doc.add_picture(image_path, width=w, height=h)
    save_document(doc, path)
    return "Image added."

def add_picture_stream(path: str, image_base64: str, width_inches: Optional[float] = None) -> str:
    """Add an image from Base64 string."""
    doc = open_document(path)
    image_data = base64.b64decode(image_base64)
    image_stream = io.BytesIO(image_data)
    w = Inches(width_inches) if width_inches else None
    doc.add_picture(image_stream, width=w)
    save_document(doc, path)
    return "Image stream added."
