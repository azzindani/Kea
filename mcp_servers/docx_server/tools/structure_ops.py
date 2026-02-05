from mcp_servers.docx_server.tools.core_ops import open_document, save_document
import docx
from docx.enum.section import WD_SECTION
from typing import List, Dict, Any, Optional

def add_section(path: str, section_type: str = "NEXT_PAGE") -> str:
    """Add a new section."""
    doc = open_document(path)
    enum_type = getattr(WD_SECTION, section_type, WD_SECTION.NEXT_PAGE)
    doc.add_section(enum_type)
    save_document(doc, path)
    return "Section added."

def get_sections(path: str) -> List[Dict[str, Any]]:
    """List sections info."""
    doc = open_document(path)
    info = []
    for i, section in enumerate(doc.sections):
        info.append({
            "index": i,
            "start_type": str(section.start_type),
            "page_height": section.page_height.inches if section.page_height else None,
            "page_width": section.page_width.inches if section.page_width else None
        })
    return info

def set_header(path: str, section_index: int, text: str) -> str:
    """Set header text for a section."""
    doc = open_document(path)
    if section_index >= len(doc.sections): raise IndexError("Section index out of range")
    section = doc.sections[section_index]
    section.header.paragraphs[0].text = text
    save_document(doc, path)
    return "Header updated."

def set_footer(path: str, section_index: int, text: str) -> str:
    """Set footer text for a section."""
    doc = open_document(path)
    if section_index >= len(doc.sections): raise IndexError("Section index out of range")
    section = doc.sections[section_index]
    section.footer.paragraphs[0].text = text
    save_document(doc, path)
    return "Footer updated."

def read_header(path: str, section_index: int) -> str:
    """Read header text."""
    doc = open_document(path)
    if section_index >= len(doc.sections): return ""
    return "\n".join([p.text for p in doc.sections[section_index].header.paragraphs])

def read_footer(path: str, section_index: int) -> str:
    """Read footer text."""
    doc = open_document(path)
    if section_index >= len(doc.sections): return ""
    return "\n".join([p.text for p in doc.sections[section_index].footer.paragraphs])

def add_page_break(path: str) -> str:
    """Add a hard page break."""
    doc = open_document(path)
    doc.add_page_break()
    save_document(doc, path)
    return "Page break added."
